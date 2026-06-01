import docx
import sys

def read_docx(file_path):
    doc = docx.Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    
    # Also check tables
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text for cell in row.cells]
            full_text.append(" | ".join(row_text))
            
    return "\n".join(full_text)

if __name__ == "__main__":
    path = r"c:\Users\ay024\OneDrive\Desktop\DT_projects\sparsh2.0\ORM_Dev_Working 1 (1).docx"
    content = read_docx(path)
    with open("scratch/docx_content.txt", "w", encoding="utf-8") as f:
        f.write(content)
    print("Content written to scratch/docx_content.txt")
